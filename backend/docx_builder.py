# ==========================================================
# PricePilot AI - Enterprise DOCX Document Generator Helper
# Creates Microsoft Word (.docx) documents matching PDF standards
# ==========================================================

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

NAVY_BLUE = RGBColor(0x1E, 0x3A, 0x8A)
INDIGO = RGBColor(0x43, 0x38, 0xCA)
DARK_GRAY = RGBColor(0x1F, 0x29, 0x37)
SLATE_GRAY = RGBColor(0x47, 0x55, 0x69)
LIGHT_BG = "F8FAFC"
BORDER_COLOR = "CBD5E1"

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_docx_report(filepath, doc_title, subtitle, metadata_rows, sections):
    doc = Document()
    
    # Page Setup - Standard Letter with 0.75 inch margins
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    # 1. Cover Page Header
    org_p = doc.add_paragraph()
    org_run = org_p.add_run("INFOSYS SPRINGBOARD 7.0 — PRICEPILOT AI ENTERPRISE PLATFORM")
    org_run.font.name = "Arial"
    org_run.font.size = Pt(9)
    org_run.font.bold = True
    org_run.font.color.rgb = INDIGO
    org_p.paragraph_format.space_after = Pt(6)

    # Document Title
    t_p = doc.add_paragraph()
    t_run = t_p.add_run(doc_title)
    t_run.font.name = "Arial"
    t_run.font.size = Pt(22)
    t_run.font.bold = True
    t_run.font.color.rgb = NAVY_BLUE
    t_p.paragraph_format.space_after = Pt(4)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(subtitle)
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = SLATE_GRAY
    sub_p.paragraph_format.space_after = Pt(14)

    # Metadata Table
    meta_tbl = doc.add_table(rows=len(metadata_rows), cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.autofit = False

    for idx, (k, v) in enumerate(metadata_rows):
        row_cells = meta_tbl.rows[idx].cells
        row_cells[0].width = Inches(2.0)
        row_cells[1].width = Inches(5.0)

        p0 = row_cells[0].paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.name = "Arial"
        r0.font.size = Pt(9)
        r0.font.bold = True
        r0.font.color.rgb = DARK_GRAY

        p1 = row_cells[1].paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.name = "Arial"
        r1.font.size = Pt(9)
        r1.font.color.rgb = DARK_GRAY

        set_cell_background(row_cells[0], LIGHT_BG)
        set_cell_background(row_cells[1], LIGHT_BG)
        set_cell_margins(row_cells[0])
        set_cell_margins(row_cells[1])

    doc.add_page_break()

    # 2. Add Sections Content
    for sec in sections:
        stype = sec.get("type", "heading")
        text = sec.get("text", "")

        if stype == "h1":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = NAVY_BLUE
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)

        elif stype == "h2":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(11.5)
            r.font.bold = True
            r.font.color.rgb = INDIGO
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)

        elif stype == "h3":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = DARK_GRAY
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)

        elif stype == "paragraph":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(9.5)
            r.font.color.rgb = DARK_GRAY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15

        elif stype == "bullet":
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(9.5)
            r.font.color.rgb = DARK_GRAY
            p.paragraph_format.space_after = Pt(3)

        elif stype == "code":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
            r.font.color.rgb = DARK_GRAY
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)

        elif stype == "table":
            headers = sec.get("headers", [])
            rows_data = sec.get("data", [])

            tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header Row
            hdr_cells = tbl.rows[0].cells
            for c_idx, h_text in enumerate(headers):
                p = hdr_cells[c_idx].paragraphs[0]
                r = p.add_run(str(h_text))
                r.font.name = "Arial"
                r.font.size = Pt(9)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_background(hdr_cells[c_idx], "1E3A8A")
                set_cell_margins(hdr_cells[c_idx])

            # Data Rows
            for r_idx, row in enumerate(rows_data):
                row_cells = tbl.rows[r_idx + 1].cells
                bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
                for c_idx, val in enumerate(row):
                    p = row_cells[c_idx].paragraphs[0]
                    r = p.add_run(str(val))
                    r.font.name = "Arial"
                    r.font.size = Pt(8.5)
                    r.font.color.rgb = DARK_GRAY
                    set_cell_background(row_cells[c_idx], bg_color)
                    set_cell_margins(row_cells[c_idx])

            p_sp = doc.add_paragraph()
            p_sp.paragraph_format.space_after = Pt(6)

    doc.save(filepath)
    print(f"  [DOCX] Saved {os.path.basename(filepath)} ({os.path.getsize(filepath):,} bytes)")
    return filepath
